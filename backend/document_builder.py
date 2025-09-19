from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn



def crear_doc(template_file=None):
    if template_file:
        return Document(template_file.file)
    else:
        return Document("templates/report-template.docx")

def insertar_memoria_descriptiva(doc, data):
    """
    Inserta la sección de Memoria Descriptiva con el Autor, Objeto y Emplazamiento.
    """
    doc.add_heading("MEMORIA DESCRIPTIVA", level=1)
    
    # 1. AUTOR
    doc.add_heading("1. Autor", level=2)
    doc.add_paragraph(f"El presente informe ha sido redactado por {data.get('arquitecta','')}, arquitecta y colegiada con el nº 671029 del C.O.A.V.N. (Colegio Oficial de Arquitectos Vasco Navarro). El trabajo ha sido realizado en relación a la oferta aceptada '033/2023', presentada el pasado 2 de marzo del 2021.")
    
    # 2. OBJETO
    doc.add_heading("2. Objeto", level=2)
    doc.add_paragraph("Tiene la presente memoria como finalidad comunicar a la comunidad de vecinos el estado en el que se encuentra su edificio indicando en el informe, las posibles causas y posible solución. Para ello, se realiza una inspección sanitaria general, no exhaustiva.")
    
    # 3. EMPLAZAMIENTO Y COLINDANTES
    doc.add_heading("3. Emplazamiento y colindantes", level=2)
    doc.add_paragraph(f"Dirección: {data.get('direccion','')}")
    doc.add_paragraph(f"Promotor: {data.get('promotor','')}")
    doc.add_paragraph(f"Arquitecta: {data.get('arquitecta','')}")
    doc.add_paragraph(f"Fecha: {data.get('fecha','')}")
    if data.get("emplazamiento"):
        doc.add_paragraph(data["emplazamiento"])

def insertar_antecedentes(doc, data):
    if data.get("motivoInspeccion") or data.get("observacionesPrevias"):
        doc.add_heading("Antecedentes", level=1)
        if data.get("motivoInspeccion"):
            doc.add_paragraph(f"Motivo: {data['motivoInspeccion']}")
        if data.get("observacionesPrevias"):
            doc.add_paragraph(f"Observaciones: {data['observacionesPrevias']}")

def insertar_resultados(doc, data, images_map):
    doc.add_heading("RESULTADOS", level=1)
    doc.add_heading("6.1 Introducción", level=2)
    doc.add_paragraph("La inspección comienza con la visita según horarios establecidos, de las distintas viviendas que componen el edificio en cuestión. Se analiza el edificio desde la calle principal y también desde la plaza.")
    doc.add_paragraph("A continuación, se exponen las tablas con los datos obtenidos de cada una de las zonas analizadas, y comentarios acerca de lo observado en cada uno de los elementos. Se ha analizado por patologías encontradas y repeticiones de las mismas, intentando agrupar las incidencias, según los pisos afectados y vistos.")
    
    # 6.2 ESTUDIO SANITARIO
    doc.add_heading("6.2 Estudio Sanitario", level=2)
    for tipo, info in data.get("patologias", {}).items():
        if info.get("activo"):
            doc.add_heading(tipo.capitalize(), level=3)
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            table.autofit = False
            table.rows[0].cells[0].text = "Patología"
            table.rows[0].cells[0].text = tipo.capitalize()
            if info.get("pisos"):
                row = table.add_row().cells
                row[0].text = "Pisos afectados"
                row[1].text = info["pisos"]
            if info.get("descripcion"):
                row = table.add_row().cells
                row[0].text = "Descripción"
                row[1].text = info["descripcion"]                
            if info.get("grado"):
                row = table.add_row().cells
                row[0].text = "Grado de actuación"
                row[1].text = info["grado"]
                            
            # Insertar imágenes
            for img_file in images_map.get(tipo, []):
                run = doc.add_paragraph().add_run()
                run.add_picture(img_file.file, width=Inches(2))
                
def insertar_presupuesto(doc, presupuesto):
    if not presupuesto:
        return

    doc.add_heading("ANEXO IV: VALORACIÓN APROXIMADA", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    headers = ["Partida", "Unidades", "Precio Unitario", "Total"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(11)

    for item in presupuesto:
        row = table.add_row().cells
        row[0].text = item.get("partida", "")
        row[1].text = str(item.get("unidades", ""))
        row[2].text = str(item.get("precioUnitario", ""))
        row[3].text = str(item.get("total", ""))

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')
            for border_name in ("top","left","bottom","right"):
                border_el = OxmlElement(f"w:{border_name}")
                border_el.set(qn('w:val'), 'single')
                border_el.set(qn('w:sz'), '4')
                border_el.set(qn('w:space'), '0')
                border_el.set(qn('w:color'), '000000')
                tc_borders.append(border_el)
            tc_pr.append(tc_borders)

def insertar_conclusiones(doc, data):
    if data.get("conclusiones"):
        doc.add_heading("CONCLUSIONES", level=1)
        doc.add_paragraph(data["conclusiones"])

def insertar_posibles_actuaciones(doc, data):
    if data.get("posiblesActuaciones"):
        doc.add_heading("POSIBLES ACTUACIONES", level=1)
        doc.add_paragraph(data["posiblesActuaciones"])

def insertar_notas(doc, data):
    if data.get("notas"):
        doc.add_heading("NOTAS", level=1)
        doc.add_paragraph(data["notas"])


