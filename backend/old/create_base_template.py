from docx import Document
from pathlib import Path

# Ruta donde guardaremos la plantilla
output_path = Path("backend/templates/safety-report-base.docx")
output_path.parent.mkdir(parents=True, exist_ok=True)

# Crear documento nuevo
doc = Document()

# Agregar título y texto de ejemplo para asegurar que los estilos existen
doc.add_heading("Plantilla base", level=1)
doc.add_paragraph("Este documento sirve como plantilla base para generar informes Word.")

# Crear una tabla de ejemplo (para forzar que Word incluya el estilo 'Table Grid')
table = doc.add_table(rows=2, cols=2)
table.style = "Table Grid"
table.cell(0, 0).text = "Columna 1"
table.cell(0, 1).text = "Columna 2"
table.cell(1, 0).text = "Valor 1"
table.cell(1, 1).text = "Valor 2"

# Guardar documento
doc.save(output_path)
print(f"✅ Plantilla base creada en: {output_path}")
