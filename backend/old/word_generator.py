from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import json
import os

TEMPLATE_DIR = "templates"
OUTPUT_DIR = "generated"
os.makedirs(OUTPUT_DIR, exist_ok=True)

def generate_word_report(template_name: str, file_path: str, sheet_name: str, mapping: dict, output_name: str = "report.docx"):
    """
    Genera un informe Word a partir de:
      - Una plantilla Word con campos {PLACEHOLDER}
      - Un Excel con datos (hoja y columnas seleccionadas)
      - Un diccionario de mapping {columna_excel: campo_word}
    """

    # Cargar plantilla Word
    template_path = os.path.join(TEMPLATE_DIR, template_name)
    doc = Document(template_path)

    # Leer Excel
    df = pd.read_excel(file_path, sheet_name=sheet_name)

    # Reemplazo de placeholders en el documento
    for paragraph in doc.paragraphs:
        for col, placeholder in mapping.items():
            if placeholder in paragraph.text:
                value = str(df[col].iloc[0]) if col in df.columns else ""
                paragraph.text = paragraph.text.replace(placeholder, value)

    # También reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for col, placeholder in mapping.items():
                    if placeholder in cell.text:
                        value = str(df[col].iloc[0]) if col in df.columns else ""
                        cell.text = cell.text.replace(placeholder, value)

    # Guardar documento generado
    output_path = os.path.join(OUTPUT_DIR, output_name)
    doc.save(output_path)

    return output_path


def apply_paragraph_style(paragraph, style_conf: dict):
    """Aplicar estilos de párrafo desde el template JSON."""
    if "font_size" in style_conf:
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.size = Pt(style_conf["font_size"])
    if "bold" in style_conf and style_conf["bold"]:
        for run in paragraph.runs:
            run.bold = True
    if "italic" in style_conf and style_conf["italic"]:
        for run in paragraph.runs:
            run.italic = True
    if "word_style" in style_conf:
        paragraph.style = style_conf["word_style"]


def add_table(document, data: list, columns: list, style_conf: dict):
    """Insertar tabla en Word desde lista de dicts."""
    if not data:
        return

    table = document.add_table(rows=1, cols=len(columns))
    hdr_cells = table.rows[0].cells

    # Cabeceras
    for i, col in enumerate(columns):
        hdr_cells[i].text = col["label"]

    # Filas
    for row in data:
        row_cells = table.add_row().cells
        for i, col in enumerate(columns):
            row_cells[i].text = str(row.get(col["key"], ""))

    # Aplicar estilo
    if style_conf and "word_style" in style_conf:
        table.style = style_conf["word_style"]

    return table


def generate_report_from_excel(excel_path: str, template_def: dict, output_path: str):
    """
    Genera un documento Word a partir de un Excel y un template JSON.
    - excel_path: Excel temporal con los datos ya transformados
    - template_def: dict con estructura del documento
    - output_path: destino .docx
    """
    # Abrir documento base con estilos
    base_doc_path = template_def.get("base_docx", None)
    if base_doc_path:
        document = Document(base_doc_path)
    else:
        document = Document()

    # Leer Excel
    xls = pd.ExcelFile(excel_path)

    for section in template_def["structure"]:
        document.add_heading(section["title"], level=1)

        for field in section["fields"]:
            fkey = field["key"]
            ftype = field["type"]

            # Campo tipo texto / richtext
            if ftype in ("text", "richtext"):
                p = document.add_paragraph(f"[{fkey}] (rellenar desde form)")
                style_id = field.get("style", "normal")
                style_conf = template_def.get("styles", {}).get(style_id, {})
                apply_paragraph_style(p, style_conf)

            # Campo tipo tabla (leer de Excel)
            elif ftype == "table":
                if fkey in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=fkey)
                    data = df.to_dict(orient="records")
                    table_style = template_def.get("styles", {}).get("table_caption", {})
                    add_table(document, data, field["columns"], table_style)
                else:
                    document.add_paragraph(f"(No data for {fkey})")

    # Guardar Word
    document.save(output_path)
    print(f"✅ Documento generado: {output_path}")
