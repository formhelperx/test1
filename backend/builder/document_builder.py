from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

def crear_doc(template_file=None):
    if template_file:
        return Document(template_file.file)
    else:
        return Document("templates/report-template.docx")
def current_date_format(date):
    months = ("Enero", "Febrero", "Marzo", "Abri", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre")
    day = date.day
    month = months[date.month - 1]
    year = date.year
    messsage = "{} de {} del {}".format(day, month, year)

    return messsage
    
def insertar_portada(doc, data, fotoPrincipal):
    
    datos_informe = {
        '[[TITULO]]': data.get('titulo'),
        '[[DIRECCION]]': data.get('direccion'),
        '[[PROMOTOR]]': data.get('promotor'),
        '[[ARQUITECTA]]': data.get('arquitecta'),
        '[[FECHA]]': current_date_format(date.today())
    }
    for p in doc.paragraphs:
    # Usamos una copia del texto para evitar problemas de iteración
        texto_original = p.text
        for marcador, valor in datos_informe.items():
            if marcador in texto_original:
                p.text = p.text.replace(marcador, str( valor))

    #  Insertar la foto de portada
    if fotoPrincipal:
        for p in doc.paragraphs:
            if '[[FOTO_PORTADA]]' in p.text:
                p.clear()  # Borra el marcador de texto                
                p.add_run().add_picture(fotoPrincipal.file, width=Inches(5)) 
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                break 
    

def insertar_memoria_descriptiva(doc, data, fotoEmplazamiento):
    """
    Inserta la sección de Memoria Descriptiva con el Autor, Objeto y Emplazamiento.
    """
    doc.add_heading("MEMORIA DESCRIPTIVA", level=1)
    
    # 1. AUTOR
    doc.add_heading("Autor", level=2)
    doc.add_paragraph(f"El presente informe ha sido redactado por {data.get('arquitecta','')}, arquitecta y colegiada con el nº {data.get('nColegiada','')} del C.O.A.V.N. (Colegio Oficial de Arquitectos Vasco Navarro). El trabajo ha sido realizado en relación a la oferta aceptada {data.get('nOferta','')}, presentada el pasado 2 de marzo del 2021.")
    
    # 2. OBJETO
    doc.add_heading("Objeto", level=2)
    doc.add_paragraph("Tiene la presente memoria como finalidad comunicar a la comunidad de vecinos el estado en el que se encuentra su edificio indicando en el informe, las posibles causas y posible solución. Para ello, se realiza una inspección sanitaria general, no exhaustiva.")
    
    # 3. EMPLAZAMIENTO Y COLINDANTES
    doc.add_heading("Emplazamiento y colindantes", level=2)
    doc.add_paragraph (f"La inspección a realizar se ha llevado a cabo en un edificio residencial ubicado en {data.get('direccion','')}. ")
    doc.add_paragraph (f"El inmueble presenta una tipología constructiva con estructura de {data.get('tipologiaEstructura', '')}, cerramientos de {data.get('cerramientos','')} y revestimiento exterior de {data.get('revestimientoExterior','')}.")
    doc.add_paragraph (f"La cubierta es {data.get('cubiertaTipo','')}.")
    if data.get('patiosDescripcion',''):
        doc.add_paragraph (data.get('patiosDescripcion',''))

        

    if fotoEmplazamiento:
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(fotoEmplazamiento.file, width=Inches(4))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Emplazamiento", style="Figura numerada")

def insertar_antecedentes(doc, data):
    if data.get("motivoInspeccion") or data.get("observacionesPrevias"):
        doc.add_heading("Antecedentes", level=1)
        if data.get("motivoInspeccion"):
            doc.add_paragraph(f"Motivo: {data['motivoInspeccion']}")
        if data.get("observacionesPrevias"):
            doc.add_paragraph(f"Observaciones: {data['observacionesPrevias']}")

def insertar_resultados(doc, data, images_map):
    doc.add_heading("RESULTADOS", level=1)
    doc.add_heading("Introducción", level=2)
    doc.add_paragraph(
        "La inspección comienza con la visita según horarios establecidos, "
        "de las distintas viviendas que componen el edificio en cuestión. "
        "Se analiza el edificio desde la calle principal y también desde la plaza."
    )
    doc.add_paragraph(
        "A continuación, se exponen las tablas con los datos obtenidos de cada una de las zonas analizadas, "
        "y comentarios acerca de lo observado en cada uno de los elementos. "
        "Se ha analizado por patologías encontradas y repeticiones de las mismas, "
        "intentando agrupar las incidencias según los pisos afectados y vistos."
    )   
    # 6.2 ESTUDIO SANITARIO
    for idx, pat in enumerate(data.get("patologias", [])):
        doc.add_heading(pat.get("tipo", "Patología sin nombre"), level=3)

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.autofit = False

        row = table.rows[0].cells
        row[0].text = "Tipo de Patología"
        row[1].text = pat.get("tipo", "")

        if pat.get("pisos"):
            row = table.add_row().cells
            row[0].text = "Pisos afectados"
            row[1].text = pat["pisos"]

        if pat.get("descripcion"):
            row = table.add_row().cells
            row[0].text = "Descripción"
            row[1].text = pat["descripcion"]

        if pat.get("grado"):
            row = table.add_row().cells
            row[0].text = "Grado de actuación"
            row[1].text = str(pat["grado"])

        # Insertar imágenes si existen
        if idx in images_map:
            for img_file in images_map[idx]:
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_picture(img_file.file, width=Inches(3))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph(pat.get("tipo", ""), style="Figura numerada")
                
def insertar_presupuesto(doc, presupuesto):
    if not presupuesto:
        return

    doc.add_heading("ANEXO: VALORACIÓN APROXIMADA", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    headers = ["Partida", "Unidades", "Precio Unitario", "Total"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(11)

    for item in presupuesto:
        row = table.add_row().cells
        row[0].text = item.get("partida", "")
        row[1].text = str(item.get("unidades", ""))
        row[2].text = str(item.get("precioUnitario", ""))
        row[3].text = str(item.get("total", ""))

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')
            for border_name in ("top","left","bottom","right"):
                border_el = OxmlElement(f"w:{border_name}")
                border_el.set(qn('w:val'), 'single')
                border_el.set(qn('w:sz'), '4')
                border_el.set(qn('w:space'), '0')
                border_el.set(qn('w:color'), '000000')
                tc_borders.append(border_el)
            tc_pr.append(tc_borders)

def insertar_conclusiones(doc, data):
    if data.get("conclusiones"):
        doc.add_heading("CONCLUSIONES", level=1)
        doc.add_paragraph(data["conclusiones"])

def insertar_posibles_actuaciones(doc, data):
    if data.get("posiblesActuaciones"):
        doc.add_heading("POSIBLES ACTUACIONES", level=1)
        doc.add_paragraph(data["posiblesActuaciones"])

def insertar_notas(doc, data):
    if data.get("notas"):
        doc.add_heading("NOTAS", level=1)
        doc.add_paragraph(data["notas"])


